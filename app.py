from flask import Flask, request, send_file, jsonify, render_template
import os
from docx import Document
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF


app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files.get('file')
    if file and file.filename.endswith('.docx'):
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        return {"message": "File uploaded successfully", "filepath": filepath}, 200
    return {"error": "Invalid file type"}, 400

@app.route('/metadata', methods=['GET'])
def metadata():
    filepath = request.args.get('filepath')
    if os.path.exists(filepath):
        doc = Document(filepath)
        core_properties = doc.core_properties
        metadata = {
            "author": core_properties.author or "Unknown",
            "last_modified": core_properties.modified or "Unknown",
            "title": core_properties.title or "No Title",
        }
        return jsonify(metadata), 200
    return {"error": "File not found"}, 404

@app.route('/convert', methods=['POST'])
def convert_to_pdf():
    filepath = request.json.get('filepath')
    if os.path.exists(filepath):
        pdf_path = filepath.replace('.docx', '.pdf')
        doc = Document(filepath)

        # Initialize PDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Add Unicode-compatible font
        font_path = "DejaVuSans.ttf"  # Adjust path if needed
        pdf.add_font("DejaVu", style="", fname=font_path, uni=True)
        pdf.set_font("DejaVu", size=12)

        # Add content from the docx file
        for paragraph in doc.paragraphs:
            if len(paragraph.text.strip()) > 0:
                pdf.multi_cell(0, 10, paragraph.text)

        # Save the PDF
        pdf.output(pdf_path, "F")
        return send_file(pdf_path, as_attachment=True), 200
    return {"error": "File not found"}, 404

@app.route('/add_password', methods=['POST'])
def add_password():
    filepath = request.json.get('filepath')
    password = request.json.get('password')
    if os.path.exists(filepath):
        pdf_path = filepath.replace('.docx', '.pdf')
        protected_pdf_path = pdf_path.replace(".pdf", "_protected.pdf")
        if os.path.exists(pdf_path):
            writer = PdfWriter()
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)
            writer.encrypt(password)
            with open(protected_pdf_path, 'wb') as f:
                writer.write(f)
            return send_file(protected_pdf_path, as_attachment=True), 200
        return {"error": "PDF file not found to protect"}, 404
    return {"error": "File not found"}, 404

if __name__ == '__main__':
    app.run(debug=True)

